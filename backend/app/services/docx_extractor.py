"""
DOCX text extraction — deterministic, no LLM.

Uses python-docx to pull plain text from uploaded Word documents.
Handles paragraphs, tables, and headers — all common in resumes.
"""

import io
import logging

import docx

from app.services.security import validate_extracted_text

logger = logging.getLogger(__name__)


class DOCXExtractionError(Exception):
    """Raised when DOCX text extraction fails or yields no usable text."""
    pass


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract plain text from a DOCX file.

    Reads paragraphs, tables, and headers. Resumes commonly use all three,
    especially tables for two-column layouts.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Concatenated text content, whitespace-normalized.

    Raises:
        DOCXExtractionError: If the file can't be read or yields no text.
    """
    if not file_bytes:
        raise DOCXExtractionError("Received empty file bytes.")

    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise DOCXExtractionError(
            f"Failed to open file as DOCX: {e}. "
            "The file may be corrupted or not a valid Word document."
        ) from e

    parts: list[str] = []

    # Extract paragraph text (covers headings, body text, bullet points)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Extract table text (resumes often use tables for layout)
    for table in document.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" | ".join(row_texts))

    if not parts:
        raise DOCXExtractionError(
            "DOCX was read successfully but contained no extractable text. "
            "The file may be empty or contain only images."
        )

    full_text = "\n".join(parts)

    # Validate minimum content
    try:
        full_text = validate_extracted_text(full_text, "DOCX document")
    except Exception as e:
        raise DOCXExtractionError(str(e)) from e

    logger.info(
        "Extracted %d characters from DOCX (%d paragraphs, %d tables)",
        len(full_text),
        len(document.paragraphs),
        len(document.tables),
    )
    return full_text
