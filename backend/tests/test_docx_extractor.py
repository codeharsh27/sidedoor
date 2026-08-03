"""
Tests for DOCX text extraction.

Creates DOCX files in memory using python-docx so no fixture files are needed.
Tests cover paragraphs, tables, empty docs, and non-DOCX input.
"""

import io

import docx
import pytest

from app.services.docx_extractor import DOCXExtractionError, extract_text_from_docx


def _make_docx_with_paragraphs(texts: list[str]) -> bytes:
    """Create a DOCX with the given paragraphs and return as bytes."""
    document = docx.Document()
    for text in texts:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_docx_with_table(rows: list[list[str]]) -> bytes:
    """Create a DOCX with a table and return as bytes."""
    document = docx.Document()
    document.add_paragraph("Resume")
    table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestExtractTextFromDocx:
    """Tests for extract_text_from_docx()."""

    def test_extracts_paragraphs(self):
        """Should extract text from paragraphs."""
        docx_bytes = _make_docx_with_paragraphs([
            "John Doe",
            "Software Engineer",
            "Skills: Python, React, PostgreSQL, Docker",
            "Experience building REST APIs and data pipelines.",
        ])
        result = extract_text_from_docx(docx_bytes)

        assert "John Doe" in result
        assert "Python" in result
        assert "REST APIs" in result

    def test_extracts_table_content(self):
        """Should extract text from tables (common in resume layouts)."""
        docx_bytes = _make_docx_with_table([
            ["Skill", "Experience"],
            ["Python", "3 years"],
            ["React", "2 years"],
        ])
        result = extract_text_from_docx(docx_bytes)

        assert "Python" in result
        assert "3 years" in result

    def test_combines_paragraphs_and_tables(self):
        """Should extract from both paragraphs and tables."""
        document = docx.Document()
        document.add_paragraph("John Doe - Software Engineer with extensive experience in backend systems")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Language"
        table.cell(0, 1).text = "Level"
        table.cell(1, 0).text = "Python"
        table.cell(1, 1).text = "Expert"
        document.add_paragraph("Built multiple production REST APIs serving thousands of requests per day")

        buffer = io.BytesIO()
        document.save(buffer)
        docx_bytes = buffer.getvalue()

        result = extract_text_from_docx(docx_bytes)
        assert "John Doe" in result
        assert "Python" in result
        assert "REST APIs" in result

    def test_raises_on_empty_bytes(self):
        """Empty bytes should raise DOCXExtractionError."""
        with pytest.raises(DOCXExtractionError, match="empty"):
            extract_text_from_docx(b"")

    def test_raises_on_invalid_bytes(self):
        """Non-DOCX bytes should raise DOCXExtractionError."""
        with pytest.raises(DOCXExtractionError):
            extract_text_from_docx(b"this is not a docx file at all")

    def test_raises_on_empty_document(self):
        """A DOCX with no text content should raise DOCXExtractionError."""
        document = docx.Document()
        # Add only empty paragraphs
        document.add_paragraph("")
        document.add_paragraph("   ")
        buffer = io.BytesIO()
        document.save(buffer)

        with pytest.raises(DOCXExtractionError):
            extract_text_from_docx(buffer.getvalue())

    def test_returns_string(self):
        """Result should be a string."""
        docx_bytes = _make_docx_with_paragraphs([
            "Python developer with experience in building web applications and APIs",
        ])
        result = extract_text_from_docx(docx_bytes)
        assert isinstance(result, str)

    def test_strips_empty_paragraphs(self):
        """Empty paragraphs should not produce blank lines in output."""
        docx_bytes = _make_docx_with_paragraphs([
            "Content line one with enough text to pass minimum threshold checks",
            "",
            "",
            "Content line two with more relevant resume information for testing",
        ])
        result = extract_text_from_docx(docx_bytes)
        # Should not have triple newlines from empty paragraphs
        assert "\n\n\n" not in result
