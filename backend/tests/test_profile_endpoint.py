"""
Integration tests for the profile parse endpoints.

Uses FastAPI TestClient with mocked LLM, embedder, and DB session.
Verifies the full request → extract → parse → embed → response flow
for all three input paths: file upload, raw text, and portfolio URL.
"""

import io
import json
import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import docx
import pytest
from fastapi.testclient import TestClient

from app.services.resume_parser import NotableProject, ProfileData

# Sample parsed profile for mocking
MOCK_PROFILE = ProfileData(
    skills=["Python", "FastAPI", "PostgreSQL"],
    domains=["backend web dev"],
    project_summary="Junior backend developer with 1 internship.",
    notable_projects=[
        NotableProject(
            title="Task API",
            description="REST API for task management with JWT auth.",
            tech_used=["Python", "FastAPI", "SQLAlchemy"],
        )
    ],
)


def _make_test_docx() -> bytes:
    """Create a simple DOCX file for testing."""
    document = docx.Document()
    document.add_paragraph("John Doe - Software Engineer with extensive experience in Python and backend systems")
    document.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker, Kubernetes, Redis")
    document.add_paragraph("Experience: 3 years building REST APIs and microservices at scale")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def mock_parser():
    """A resume parser that returns MOCK_PROFILE without calling any LLM."""
    parser = MagicMock()
    parser.parse_resume = AsyncMock(return_value=MOCK_PROFILE)
    return parser


@pytest.fixture
def mock_embedder():
    """An embedder that returns a fixed 384-dim vector without loading a model."""
    embedder = MagicMock()
    embedder.embed_profile.return_value = [0.1] * 384
    embedder.dimensions = 384
    return embedder


@pytest.fixture
def test_user_id():
    """A fixed UUID for the test user."""
    return uuid.UUID("12345678-1234-5678-1234-567812345678")


@pytest.fixture
def mock_db_session(test_user_id):
    """
    A mocked async DB session.

    Simulates:
    - User exists (for FK check)
    - No existing profile (first upload)
    """
    session = AsyncMock()

    # Mock user lookup: return a user object
    mock_user = MagicMock()
    mock_user.id = test_user_id

    # First execute() call returns user (for the user lookup)
    # Second execute() call returns None (no existing profile)
    user_result = MagicMock()
    user_result.scalar_one_or_none.return_value = mock_user

    profile_result = MagicMock()
    profile_result.scalar_one_or_none.return_value = None

    session.execute = AsyncMock(side_effect=[user_result, profile_result])
    session.add = MagicMock()
    session.commit = AsyncMock()
    session.rollback = AsyncMock()

    return session


@pytest.fixture
def client(mock_parser, mock_embedder, mock_db_session):
    """FastAPI test client with all dependencies mocked."""
    from app.api.routes.profile import get_resume_parser
    from app.db.session import get_db_session

    with patch("app.services.embedder.SentenceTransformer") as mock_transformer_cls:
        # Mock the dimension method so loading the model doesn't crash/block
        mock_transformer_cls.return_value.get_sentence_embedding_dimension.return_value = 384
        
        from app.main import app

        async def override_db():
            yield mock_db_session

        app.dependency_overrides[get_db_session] = override_db
        app.dependency_overrides[get_resume_parser] = lambda: mock_parser

        with patch("app.api.routes.profile.get_embedder", return_value=mock_embedder):
            with TestClient(app, raise_server_exceptions=False) as c:
                yield c

        app.dependency_overrides.clear()


# ---------- POST /api/v1/profile/parse-text ----------


class TestParseTextEndpoint:
    """Tests for POST /api/v1/profile/parse-text."""

    def test_successful_parse(self, client, test_user_id):
        """Valid request should return parsed profile data."""
        response = client.post(
            "/api/v1/profile/parse-text",
            json={
                "user_id": str(test_user_id),
                "raw_text": "John Doe, Python developer with 1 year of experience in web development and backend systems.",
            },
        )

        assert response.status_code == 200
        data = response.json()
        assert data["user_id"] == str(test_user_id)
        assert "Python" in data["skills"]
        assert data["source_type"] == "text"
        assert len(data["notable_projects"]) == 1
        assert data["notable_projects"][0]["title"] == "Task API"

    def test_empty_text_rejected(self, client, test_user_id):
        """Empty resume text should return 400."""
        response = client.post(
            "/api/v1/profile/parse-text",
            json={
                "user_id": str(test_user_id),
                "raw_text": "",
            },
        )
        assert response.status_code == 400

    def test_too_short_text_rejected(self, client, test_user_id):
        """Text below minimum threshold should return 400."""
        response = client.post(
            "/api/v1/profile/parse-text",
            json={
                "user_id": str(test_user_id),
                "raw_text": "too short",
            },
        )
        assert response.status_code == 400

    def test_invalid_user_id_rejected(self, client):
        """Non-UUID user_id should return 400."""
        response = client.post(
            "/api/v1/profile/parse-text",
            json={
                "user_id": "not-a-uuid",
                "raw_text": "Some resume text that is long enough to pass the minimum threshold check.",
            },
        )
        assert response.status_code == 400

    def test_response_contains_source_type(self, client, test_user_id):
        """Response should include source_type field."""
        response = client.post(
            "/api/v1/profile/parse-text",
            json={
                "user_id": str(test_user_id),
                "raw_text": "Python developer with experience in web services and API development for 3 years.",
            },
        )

        assert response.status_code == 200
        data = response.json()
        assert "source_type" in data
        assert data["source_type"] == "text"

    def test_response_shape(self, client, test_user_id):
        """Response should contain all expected fields."""
        response = client.post(
            "/api/v1/profile/parse-text",
            json={
                "user_id": str(test_user_id),
                "raw_text": "Resume content here with enough text to pass the minimum validation threshold.",
            },
        )

        assert response.status_code == 200
        data = response.json()

        # All fields present
        assert "user_id" in data
        assert "skills" in data
        assert "domains" in data
        assert "project_summary" in data
        assert "notable_projects" in data
        assert "source_type" in data

        # Types correct
        assert isinstance(data["skills"], list)
        assert isinstance(data["domains"], list)
        assert isinstance(data["project_summary"], str)
        assert isinstance(data["notable_projects"], list)
        assert isinstance(data["source_type"], str)


# ---------- POST /api/v1/profile/parse (file upload) ----------


class TestParseFileEndpoint:
    """Tests for POST /api/v1/profile/parse with file uploads."""

    def test_unsupported_file_type_rejected(self, client, test_user_id):
        """Non-PDF/DOCX files should be rejected."""
        response = client.post(
            "/api/v1/profile/parse",
            data={"user_id": str(test_user_id)},
            files={"file": ("image.png", b"fake png content", "image/png")},
        )
        assert response.status_code == 422

    def test_empty_file_rejected(self, client, test_user_id):
        """Empty files should be rejected."""
        response = client.post(
            "/api/v1/profile/parse",
            data={"user_id": str(test_user_id)},
            files={"file": ("resume.pdf", b"", "application/pdf")},
        )
        assert response.status_code == 422

    def test_docx_upload_accepted(self, client, test_user_id):
        """DOCX files should be accepted and parsed."""
        docx_bytes = _make_test_docx()
        response = client.post(
            "/api/v1/profile/parse",
            data={"user_id": str(test_user_id)},
            files={
                "file": (
                    "resume.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert response.status_code == 200
        data = response.json()
        assert data["source_type"] == "docx"
        assert "Python" in data["skills"]


# ---------- POST /api/v1/profile/parse-url ----------


class TestParseUrlEndpoint:
    """Tests for POST /api/v1/profile/parse-url."""

    def test_ssrf_private_ip_blocked(self, client, test_user_id):
        """URLs pointing to private IPs should be blocked."""
        response = client.post(
            "/api/v1/profile/parse-url",
            json={
                "user_id": str(test_user_id),
                "portfolio_url": "http://127.0.0.1/admin",
            },
        )
        # Should be 403 (blocked) or 422 (extraction error)
        assert response.status_code in (403, 422)

    def test_forbidden_scheme_blocked(self, client, test_user_id):
        """Non-HTTP schemes should be rejected."""
        response = client.post(
            "/api/v1/profile/parse-url",
            json={
                "user_id": str(test_user_id),
                "portfolio_url": "file:///etc/passwd",
            },
        )
        assert response.status_code in (403, 422)

    def test_invalid_user_id_rejected(self, client):
        """Non-UUID user_id should return 400."""
        response = client.post(
            "/api/v1/profile/parse-url",
            json={
                "user_id": "not-a-uuid",
                "portfolio_url": "https://example.com",
            },
        )
        assert response.status_code == 400


# ---------- Health check ----------


class TestHealthEndpoint:
    """Tests for GET /health."""

    def test_health_check(self, client):
        """Health endpoint should return ok."""
        response = client.get("/health")
        assert response.status_code == 200
        assert response.json() == {"status": "ok"}
